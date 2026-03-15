import io
from typing import Union

from PyPDF2 import PdfReader
import docx
import os

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF file.
    """
    pdf_stream = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_stream)

    text = ""

    for page in reader.pages:
        text += page.extract_text() or ""

    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX file.
    """
    doc_stream = io.BytesIO(file_bytes)
    document = docx.Document(doc_stream)

    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    return text


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Extract text from TXT file.
    """
    return file_bytes.decode("utf-8")


def load_document(file_bytes: bytes, filename: str) -> str:
    """
    Detect file type and extract text accordingly.
    """

    extension = filename.split(".")[-1].lower()

    if extension == "pdf":
        return extract_text_from_pdf(file_bytes)

    elif extension == "docx":
        return extract_text_from_docx(file_bytes)

    elif extension == "txt":
        return extract_text_from_txt(file_bytes)

    else:
        raise ValueError("Unsupported file type")




def load_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    elif ext == ".docx":
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    else:
        raise ValueError("Unsupported file format")